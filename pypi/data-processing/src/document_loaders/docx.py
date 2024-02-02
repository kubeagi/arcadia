# Copyright 2024 KubeAGI.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import logging
from typing import List

import docx
from langchain_core.documents import Document

from common import log_tag_const
from document_loaders.base import BaseLoader
from utils import file_utils

logger = logging.getLogger(__name__)

class DocxLoader(BaseLoader):
    """Load docx files."""

    def __init__(
        self,
        file_path: str,
    ):
        """
        Initialize the loader with a list of URL paths.

        Args:
            file_path (str): File Path.
        """
        self._file_path = file_path

    def load(self) -> List[Document]:
        """
        Load and return all Documents from the docx file.

        Returns:
            List[Document]: A list of Document objects.

        """
        logger.info(f"{log_tag_const.DOCX_LOADER} Start to load docx file")

        # Get file name
        file_name = file_utils.get_file_name(self._file_path)

        docs = []
        doc = docx.Document(self._file_path)
        for i in range(len(doc.paragraphs)):
            para = doc.paragraphs[i]
            content = para.text
            metadata = {"source": file_name, "page": i}
            docs.append(Document(page_content=content, metadata=metadata))

        return docs
